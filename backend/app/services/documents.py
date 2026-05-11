"""Document parsing utilities.

Each upload becomes a :class:`ParsedDocument` that exposes both a textual
preview (used for prompt context) and the bits we need to attach the
document to a Claude message (raw bytes + media type for PDFs and images).
DOCX files are converted to plain text because Claude does not natively
accept .docx attachments.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

import anyio
from pypdf import PdfReader

logger = logging.getLogger(__name__)


_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx", ".doc"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}

_IMAGE_MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}


@dataclass
class ParsedDocument:
    """A single uploaded document, parsed for downstream use."""

    filename: str
    """Original filename as supplied by the client."""

    kind: str
    """One of ``"pdf"``, ``"docx"``, ``"image"`` or ``"other"``."""

    raw_bytes: bytes
    """The original file bytes."""

    media_type: str = ""
    """MIME type for ``image``/``pdf`` documents (empty for text-only)."""

    text: str = ""
    """Extracted text content (best-effort)."""

    pages: int | None = None
    """Page count when known (PDFs only)."""

    error: str | None = None
    """Populated if parsing failed; the document is still kept."""

    extra: dict[str, str] = field(default_factory=dict)


def _ext_of(name: str) -> str:
    name = name.lower()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def _parse_pdf(filename: str, data: bytes) -> ParsedDocument:
    try:
        reader = PdfReader(io.BytesIO(data))
        text_parts: list[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:  # noqa: BLE001 - page-level failure
                logger.warning("PDF page %d extraction failed: %s", idx, exc)
                page_text = ""
            if page_text.strip():
                text_parts.append(f"[Page {idx}]\n{page_text.strip()}")
        return ParsedDocument(
            filename=filename,
            kind="pdf",
            raw_bytes=data,
            media_type="application/pdf",
            text="\n\n".join(text_parts),
            pages=len(reader.pages),
        )
    except Exception as exc:  # noqa: BLE001 - parsing fallback
        logger.warning("Failed to parse PDF %s: %s", filename, exc)
        return ParsedDocument(
            filename=filename,
            kind="pdf",
            raw_bytes=data,
            media_type="application/pdf",
            error=str(exc),
        )


def _parse_docx(filename: str, data: bytes) -> ParsedDocument:
    try:
        from docx import Document  # local import to keep startup fast
    except Exception as exc:  # noqa: BLE001 - dependency missing
        return ParsedDocument(
            filename=filename,
            kind="docx",
            raw_bytes=data,
            error=f"python-docx unavailable: {exc}",
        )
    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        return ParsedDocument(
            filename=filename,
            kind="docx",
            raw_bytes=data,
            text="\n".join(paragraphs),
        )
    except Exception as exc:  # noqa: BLE001 - parsing fallback
        logger.warning("Failed to parse DOCX %s: %s", filename, exc)
        return ParsedDocument(
            filename=filename,
            kind="docx",
            raw_bytes=data,
            error=str(exc),
        )


def _parse_image(filename: str, data: bytes, ext: str) -> ParsedDocument:
    media_type = _IMAGE_MEDIA_TYPES.get(ext, "image/png")
    return ParsedDocument(
        filename=filename,
        kind="image",
        raw_bytes=data,
        media_type=media_type,
    )


def _parse_one(filename: str, data: bytes, content_type: str) -> ParsedDocument:
    ext = _ext_of(filename)
    if ext in _PDF_EXTS or content_type == "application/pdf":
        return _parse_pdf(filename, data)
    if ext in _DOCX_EXTS or "wordprocessingml" in content_type or content_type == "application/msword":
        return _parse_docx(filename, data)
    if ext in _IMAGE_EXTS or content_type.startswith("image/"):
        return _parse_image(filename, data, ext)
    # Fallback: try to decode as text
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        text = ""
    return ParsedDocument(
        filename=filename,
        kind="other",
        raw_bytes=data,
        text=text,
    )


async def parse_documents(
    uploads: list[tuple[str, bytes, str]],
) -> list[ParsedDocument]:
    """Parse a list of ``(filename, bytes, content_type)`` tuples concurrently."""

    if not uploads:
        return []

    async def _run(filename: str, data: bytes, content_type: str) -> ParsedDocument:
        return await anyio.to_thread.run_sync(_parse_one, filename, data, content_type)

    results: list[ParsedDocument] = []
    async with anyio.create_task_group() as tg:
        slots: list[ParsedDocument | None] = [None] * len(uploads)

        async def _task(idx: int, filename: str, data: bytes, content_type: str) -> None:
            slots[idx] = await _run(filename, data, content_type)

        for idx, (filename, data, content_type) in enumerate(uploads):
            tg.start_soon(_task, idx, filename, data, content_type)

    for slot in slots:
        if slot is not None:
            results.append(slot)
    return results


def text_preview(doc: ParsedDocument, max_chars: int = 8000) -> str:
    """Return a trimmed text preview suitable for inclusion in a prompt."""
    if not doc.text:
        return ""
    if len(doc.text) <= max_chars:
        return doc.text
    head = doc.text[: max_chars // 2]
    tail = doc.text[-max_chars // 2 :]
    return f"{head}\n\n…[truncated]…\n\n{tail}"
